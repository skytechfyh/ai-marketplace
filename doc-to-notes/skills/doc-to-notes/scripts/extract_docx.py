#!/usr/bin/env python3
"""
Extract text structure, code blocks, lists, tables and images from .docx / .doc / .pdf
files for the doc-to-notes skill.

Usage:
    python3 extract_docx.py <doc_path> [--output-dir <dir>] [--max-img-px 2000]

Outputs:
    <output-dir>/manifest.json      — full document structure
    <output-dir>/chapter_NN.json    — one file per H2 chapter (for chunked processing)
    <output-dir>/images/            — extracted (and resized) images
    Prints: title, H2 chapter list, section-type counts, image count, paths

Key behaviours (tuned for Chinese training docs, e.g. 多易大数据):
  * CODE BLOCKS — many vendors wrap code in a single-cell (1x1) table whose first
    line is the language label (Shell / YAML / SQL / Java / Plain Text ...). These are
    emitted as fenced code blocks with the language mapped and newlines preserved,
    NOT as broken single-column markdown tables.
  * LISTS — paragraphs with numbering (numPr) keep their indent level and ordered/
    unordered type instead of collapsing into plain text.
  * HEADINGS — detected by (1) real Heading styles, (2) numbering prefix depth
    (1. → H2, 1.1 → H3, 1.1.1 → H4), (3) font-size + bold heuristics. Numbering, when
    present, overrides the level so the hierarchy is consistent.
  * IMAGES — extracted, oversized images (> max-img-px on any side) resized in place
    so Claude's vision Read calls don't hit the many-image size limit. Dimensions are
    recorded in the manifest.
  * PDF — text spans are merged per block/line (not per span) so prose isn't shredded.
  * .doc — auto-converted to .docx via macOS `textutil` when available.

Requires: python-docx (docx) ; PyMuPDF (fitz) for PDF ; Pillow optional for resize.
"""

import sys
import os
import re
import json
import shutil
import argparse
import subprocess
from pathlib import Path

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    sys.exit("[ERROR] python-docx not installed. Run: pip install python-docx")


# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

# Font-size thresholds (points) for heading detection.
# Font size is the PRIMARY, absolute signal — numbering prefixes ("1." "2.") restart
# under each section so they are NOT reliable absolute levels on their own.
H1_MIN_PT = 22.0   # document title  (observed 26pt)
H2_MIN_PT = 17.0   # top section     (observed 18pt)
H3_MIN_PT = 15.5   # sub-section     (observed 16pt)
H4_MIN_PT = 13.0   # enumerated item (observed 14-15pt)

DEFAULT_MAX_IMG_PX = 2000

# A chapter file larger than this many sections is hard to write within the 120s
# timeout, so 'auto' split drops one level deeper (H2 → H3) when chapters are this big.
MAX_CHAPTER_SECTIONS = 90

# Lower bound to avoid a swarm of tiny files: consecutive small chapters UNDER the same
# parent are merged until they reach this size (never crossing a parent boundary, never
# exceeding MAX_CHAPTER_SECTIONS). Set --min-sections 0 to disable merging.
MIN_CHAPTER_SECTIONS = 15

# Monospace fonts → paragraph is a code line (fallback when code isn't in a table)
MONOSPACE_FONTS = {
    "courier", "courier new", "consolas", "monaco", "inconsolata",
    "source code pro", "jetbrains mono", "fira code", "lucida console",
    "menlo", "andale mono", "liberation mono", "dejavu sans mono",
}

# First-line language label (lower-cased) → markdown fence language
LANG_LABEL_MAP = {
    "shell": "bash", "bash": "bash", "sh": "bash", "shellsession": "bash",
    "cmd": "bat", "powershell": "powershell", "bat": "bat",
    "yaml": "yaml", "yml": "yaml",
    "plain text": "text", "plaintext": "text", "text": "text", "txt": "text",
    "sql": "sql", "mysql": "sql", "hql": "sql", "flinksql": "sql",
    "java": "java", "scala": "scala", "kotlin": "kotlin", "groovy": "groovy",
    "python": "python", "py": "python",
    "json": "json", "xml": "xml", "html": "html", "css": "css",
    "properties": "properties", "ini": "ini", "conf": "ini",
    "markdown": "markdown", "md": "markdown",
    "dockerfile": "dockerfile", "docker": "dockerfile",
    "c": "c", "c++": "cpp", "cpp": "cpp", "go": "go", "golang": "go",
    "javascript": "javascript", "js": "javascript", "typescript": "typescript",
    "scala/java": "scala",
}

# Signals that a 1x1 table (without a known language label) is still code/config
CODE_SIGNAL_RE = re.compile(
    r"[{};=]|\b(import|package|public|private|class|def|val|var|function|"
    r"SELECT|INSERT|CREATE|FROM|WHERE|env\.|System\.|println|print\()|^\s*[#$/]"
)


# ---------------------------------------------------------------------------
# Numbering (list) resolution
# ---------------------------------------------------------------------------

class NumberingResolver:
    """Maps a paragraph's numId+ilvl to (is_ordered, level) using numbering.xml."""

    def __init__(self, doc):
        self._fmt = {}  # (numId, ilvl) -> numFmt string
        try:
            numbering = doc.part.numbering_part.numbering_definitions._numbering
        except Exception:
            return
        W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        # abstractNumId -> {ilvl: numFmt}
        abstract = {}
        for an in numbering.findall(f"{W}abstractNum"):
            aid = an.get(f"{W}abstractNumId")
            levels = {}
            for lvl in an.findall(f"{W}lvl"):
                ilvl = lvl.get(f"{W}ilvl")
                fmt_el = lvl.find(f"{W}numFmt")
                levels[ilvl] = fmt_el.get(f"{W}val") if fmt_el is not None else "bullet"
            abstract[aid] = levels
        # num -> abstractNumId
        for num in numbering.findall(f"{W}num"):
            nid = num.get(f"{W}numId")
            ref = num.find(f"{W}abstractNumId")
            aid = ref.get(f"{W}val") if ref is not None else None
            for ilvl, fmt in abstract.get(aid, {}).items():
                self._fmt[(nid, ilvl)] = fmt

    def resolve(self, para):
        """Return (is_list, is_ordered, level) for a paragraph."""
        W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        pPr = para._element.find(f"{W}pPr")
        if pPr is None:
            return (False, False, 0)
        numPr = pPr.find(f"{W}numPr")
        if numPr is None:
            return (False, False, 0)
        ilvl_el = numPr.find(f"{W}ilvl")
        numId_el = numPr.find(f"{W}numId")
        ilvl = ilvl_el.get(f"{W}val") if ilvl_el is not None else "0"
        numId = numId_el.get(f"{W}val") if numId_el is not None else None
        fmt = self._fmt.get((numId, ilvl), "bullet")
        ordered = fmt not in ("bullet", "none")
        try:
            level = int(ilvl)
        except (TypeError, ValueError):
            level = 0
        return (True, ordered, level)


# ---------------------------------------------------------------------------
# Heading detection
# ---------------------------------------------------------------------------

# e.g. "1.快速认识flink" depth 1 ; "1.1 xxx" depth 2 ; "1.1.1 xxx" depth 3
NUM_PREFIX_RE = re.compile(r"^(\d+)((?:\.\d+)*)\.?(?:\s|、|$)")


def numbering_depth(text: str):
    """Return heading depth from a leading numeric prefix, or None."""
    m = NUM_PREFIX_RE.match(text.strip())
    if not m:
        return None
    extra = m.group(2)  # ".1.2" etc
    depth = 1 + (extra.count(".") if extra else 0)
    return depth


def _max_pt(para) -> float:
    pts = [run.font.size / 12700 for run in para.runs if run.font.size]
    return max(pts) if pts else 0.0


def _is_bold(para) -> bool:
    return any(run.bold for run in para.runs if run.text.strip())


def _all_bold(para) -> bool:
    runs = [r for r in para.runs if r.text.strip()]
    return bool(runs) and all(r.bold for r in runs)


def get_heading_level(para) -> int:
    """Return 1-4 for headings, 0 for normal text. Numbering refines the level."""
    text = para.text.strip()
    if not text:
        return 0

    # 1. Real heading styles
    sname = (para.style.name if para.style else "") or ""
    if sname.startswith("Heading"):
        for p in reversed(sname.split()):
            if p.isdigit():
                return min(int(p), 4)
        return 1

    max_pt = _max_pt(para)
    bold = _is_bold(para)
    depth = numbering_depth(text)

    # 2. Explicit font size present → it is the authoritative level (absolute).
    #    Numbering prefixes are ignored here because "1./2./3." restart per section.
    if bold and max_pt > 0:
        if max_pt >= H1_MIN_PT:
            return 1
        if max_pt >= H2_MIN_PT:
            return 2
        if max_pt >= H3_MIN_PT:
            return 3
        if max_pt >= H4_MIN_PT and len(text) < 80:
            return 4
        return 0

    # 3. No explicit font size (docs that encode structure only via numbering/bold).
    if max_pt == 0 and bold:
        if depth is not None and len(text) < 80:
            # numbering depth d → heading level d+1 (H1 reserved for the title)
            return min(max(depth + 1, 2), 4)
        if _all_bold(para) and len(text) < 60:
            return 4

    return 0


# ---------------------------------------------------------------------------
# Code / cell helpers
# ---------------------------------------------------------------------------

def is_code_paragraph(para) -> bool:
    style_name = ((para.style.name if para.style else "") or "").lower()
    if "code" in style_name or "preformat" in style_name:
        return True
    for run in para.runs:
        if (run.font.name or "").lower().strip() in MONOSPACE_FONTS:
            return True
    return False


def classify_single_cell(text: str):
    """
    Classify a 1x1 table cell's content.
    Returns ('code', lang, body) | ('quote', None, text) | (None, None, text).
    """
    lines = text.split("\n")
    first = lines[0].strip().lower()
    if first in LANG_LABEL_MAP:
        lang = LANG_LABEL_MAP[first]
        body = "\n".join(lines[1:]).strip("\n")
        return ("code", lang, body)
    # No language label — decide by content shape
    if CODE_SIGNAL_RE.search(text) and ("\n" in text or len(text) < 200):
        return ("code", "text", text.strip("\n"))
    # Plain prose in a box → treat as quote/callout
    return ("quote", None, text.strip())


def cell_has_image(cell) -> bool:
    return bool(cell._element.findall(".//" + qn("a:blip")))


# ---------------------------------------------------------------------------
# Image extraction + resize
# ---------------------------------------------------------------------------

# Small-image heuristic: a low/narrow image is often a math-formula screenshot or an
# enlarged architecture-component label (e.g. "Add & Norm"), NOT a standalone figure.
# Flagged images get `small_inline:true` in the manifest so the model prioritises
# transcribing them to LaTeX ($$...$$) or folding them into prose, rather than embedding.
SMALL_INLINE_MAX_W = 700
SMALL_INLINE_MAX_H = 150


def is_small_inline(w, h) -> bool:
    return 0 < w <= SMALL_INLINE_MAX_W and 0 < h <= SMALL_INLINE_MAX_H


def image_dimensions(path: str):
    """Return (w, h) using Pillow → sips → (0,0). Best-effort."""
    try:
        from PIL import Image
        with Image.open(path) as im:
            return im.size
    except Exception:
        pass
    try:  # macOS native
        out = subprocess.check_output(
            ["sips", "-g", "pixelWidth", "-g", "pixelHeight", path],
            stderr=subprocess.DEVNULL, text=True,
        )
        w = h = 0
        for line in out.splitlines():
            if "pixelWidth" in line:
                w = int(line.split(":")[1])
            elif "pixelHeight" in line:
                h = int(line.split(":")[1])
        return (w, h)
    except Exception:
        return (0, 0)


def resize_image(path: str, max_px: int) -> tuple:
    """Resize in place if any side > max_px. Returns (w, h) after."""
    w, h = image_dimensions(path)
    if not w or not h or max(w, h) <= max_px:
        return (w, h)
    scale = max_px / max(w, h)
    nw, nh = int(w * scale), int(h * scale)
    try:
        from PIL import Image
        with Image.open(path) as im:
            im = im.convert("RGB") if im.mode in ("P", "RGBA") and path.lower().endswith((".jpg", ".jpeg")) else im
            im.resize((nw, nh), Image.LANCZOS).save(path)
        return (nw, nh)
    except Exception:
        pass
    try:  # macOS native fallback
        subprocess.check_call(
            ["sips", "--resampleHeightWidthMax", str(max_px), path],
            stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
        )
        return image_dimensions(path)
    except Exception:
        return (w, h)


def save_image(blob, ext, img_dir: Path, counter: list, max_px: int):
    """Write image bytes, resize if oversized; returns (filename, w, h)."""
    counter[0] += 1
    fname = f"image{counter[0]:03d}{ext}"
    dest = img_dir / fname
    dest.write_bytes(blob)
    w, h = resize_image(str(dest), max_px)
    return fname, w, h


# ---------------------------------------------------------------------------
# Table → markdown (real multi-column tables only)
# ---------------------------------------------------------------------------

def table_to_markdown(table) -> str:
    rows = []
    for row in table.rows:
        rows.append([c.text.strip().replace("\n", "<br/>") for c in row.cells])
    if not rows:
        return ""
    md = ["| " + " | ".join(rows[0]) + " |",
          "| " + " | ".join(["---"] * len(rows[0])) + " |"]
    for row in rows[1:]:
        md.append("| " + " | ".join(row) + " |")
    return "\n".join(md)


# ---------------------------------------------------------------------------
# DOCX extraction
# ---------------------------------------------------------------------------

def extract_docx(docx_path: str, output_dir: str, max_px: int, split_level="auto",
                 min_size=MIN_CHAPTER_SECTIONS) -> dict:
    doc = Document(docx_path)
    numbering = NumberingResolver(doc)
    img_dir = Path(output_dir) / "images"
    img_dir.mkdir(parents=True, exist_ok=True)

    img_counter = [0]
    img_seen = {}  # rid -> filename (dedupe within doc)
    sections = []

    para_map = {id(p._element): p for p in doc.paragraphs}
    table_map = {id(t._element): t for t in doc.tables}

    def emit_image_from_rid(doc, rid, caption):
        if rid in img_seen:
            fname = img_seen[rid]
            meta = next((s for s in sections if s.get("image_file") == fname), {})
            mw, mh = meta.get("width", 0), meta.get("height", 0)
            sections.append({"type": "image", "image_file": fname,
                             "caption": caption, "width": mw, "height": mh,
                             "small_inline": is_small_inline(mw, mh), "duplicate": True})
            return
        try:
            part = doc.part.related_parts[rid]
            ext = Path(part.partname).suffix or ".png"
            fname, w, h = save_image(part._blob, ext, img_dir, img_counter, max_px)
            img_seen[rid] = fname
            sections.append({"type": "image", "image_file": fname,
                             "caption": caption, "width": w, "height": h,
                             "small_inline": is_small_inline(w, h)})
        except Exception as e:
            print(f"  [WARN] image {rid}: {e}", file=sys.stderr)

    for child in doc.element.body:
        tag = child.tag.split("}")[-1]

        if tag == "p":
            para = para_map.get(id(child))
            if para is None:
                continue
            text = para.text.strip()

            # images embedded in the paragraph
            blips = para._element.findall(".//" + qn("a:blip"))
            if blips:
                for b in blips:
                    rid = b.get(qn("r:embed"))
                    if rid:
                        emit_image_from_rid(doc, rid, text or "")
                if not text:
                    continue

            if not text:
                continue

            level = get_heading_level(para)
            is_list, ordered, lvl = numbering.resolve(para)

            if level > 0:
                sections.append({"type": "heading", "level": level, "text": text})
            elif is_list:
                sections.append({"type": "list_item", "ordered": ordered,
                                 "level": lvl, "text": text})
            elif is_code_paragraph(para):
                sections.append({"type": "code", "lang": "text", "text": text})
            else:
                sections.append({"type": "paragraph", "text": text})

        elif tag == "tbl":
            table = table_map.get(id(child))
            if table is None:
                continue
            n_rows = len(table.rows)
            n_cols = len(table.columns)

            # 1x1 table → almost always a code/config block or a callout box
            if n_rows == 1 and n_cols == 1:
                cell = table.rows[0].cells[0]
                if cell_has_image(cell):
                    for b in cell._element.findall(".//" + qn("a:blip")):
                        rid = b.get(qn("r:embed"))
                        if rid:
                            emit_image_from_rid(doc, rid, cell.text.strip())
                    continue
                kind, lang, body = classify_single_cell(cell.text)
                if kind == "code":
                    sections.append({"type": "code", "lang": lang, "text": body})
                elif kind == "quote":
                    sections.append({"type": "quote", "text": body})
                continue

            md = table_to_markdown(table)
            if md:
                sections.append({"type": "table", "rows": n_rows, "cols": n_cols,
                                 "markdown": md})

    return _finalize(sections, docx_path, img_dir, img_counter[0], output_dir, split_level, min_size)


# ---------------------------------------------------------------------------
# PDF extraction (spans merged per block)
# ---------------------------------------------------------------------------

def extract_pdf(pdf_path: str, output_dir: str, max_px: int, split_level="auto",
                min_size=MIN_CHAPTER_SECTIONS) -> dict:
    try:
        import fitz
    except ImportError:
        sys.exit("[ERROR] PyMuPDF not installed. Run: pip install pymupdf")

    img_dir = Path(output_dir) / "images"
    img_dir.mkdir(parents=True, exist_ok=True)
    doc = fitz.open(pdf_path)
    sections = []
    img_counter = [0]
    xref_seen = set()

    for page_num, page in enumerate(doc):
        for img_info in page.get_images(full=True):
            xref = img_info[0]
            if xref in xref_seen:
                continue
            xref_seen.add(xref)
            try:
                base = doc.extract_image(xref)
                fname, w, h = save_image(base["image"], "." + base["ext"],
                                         img_dir, img_counter, max_px)
                sections.append({"type": "image", "image_file": fname,
                                 "caption": f"Page {page_num+1}", "width": w, "height": h,
                                 "small_inline": is_small_inline(w, h)})
            except Exception as e:
                print(f"  [WARN] pdf image p{page_num+1}: {e}", file=sys.stderr)

        for block in page.get_text("dict")["blocks"]:
            if block["type"] != 0:
                continue
            # Merge all spans in the block into one text + track max font size
            block_text_lines = []
            max_size = 0.0
            any_bold = False
            for line in block["lines"]:
                parts = []
                for span in line["spans"]:
                    t = span["text"]
                    if t:
                        parts.append(t)
                        max_size = max(max_size, span["size"])
                        if span["flags"] & 16:
                            any_bold = True
                if parts:
                    block_text_lines.append("".join(parts))
            text = "\n".join(block_text_lines).strip()
            if not text:
                continue

            depth = numbering_depth(text.split("\n")[0])
            if any_bold and max_size >= H1_MIN_PT:
                level = 1
            elif any_bold and max_size >= H2_MIN_PT:
                level = 2
            elif any_bold and max_size >= H3_MIN_PT:
                level = 3
            elif depth is not None and len(text) < 80:
                level = min(max(depth + 1, 2), 4)
            else:
                level = 0

            if level > 0:
                sections.append({"type": "heading", "level": level,
                                 "text": text.replace("\n", " ")})
            else:
                sections.append({"type": "paragraph", "text": text.replace("\n", " ")})

    return _finalize(sections, pdf_path, img_dir, img_counter[0], output_dir, split_level, min_size)


# ---------------------------------------------------------------------------
# Finalize: title, chapter split, write manifest + per-chapter JSON
# ---------------------------------------------------------------------------

def _headings_at(sections, level, title):
    return [s["text"] for s in sections
            if s.get("type") == "heading" and s.get("level") == level and s["text"] != title]


def _max_chunk_size(sections, level, title) -> int:
    """Largest number of sections between consecutive headings of `level`."""
    sizes, cur = [], 0
    seen = False
    for s in sections:
        if s.get("type") == "heading" and s.get("level") == level and s["text"] != title:
            if seen:
                sizes.append(cur)
            cur, seen = 0, True
        cur += 1
    if seen:
        sizes.append(cur)
    return max(sizes) if sizes else len(sections)


def _pick_split_level(sections, title, requested) -> int:
    """Choose the heading level to split chapter files at.

    0 (--no-split): single chapter_01.json with ALL content.
    Default (auto): split at TOP-LEVEL chapters (H2) — one file per 一级章节 — keeping the
    file count low and natural (a doc's `1.` `2.` `3.` headings are the obvious units).
    Falls back to H1, then H3, only when H2 is absent.

    A single very large H2 is deliberately NOT auto-split into many files: the 120s timeout
    limits one Edit, not one file, so a big chapter is written safely via skeleton +
    per-section fill. If you do want a huge chapter broken up, pass --split-level 3.
    """
    if requested == 0:
        return 0  # no-split mode: everything in one file
    if requested in (1, 2, 3, 4):
        return requested
    for lvl in (2, 1, 3):
        if len(_headings_at(sections, lvl, title)) >= 2:
            return lvl
    return 2


def _finalize(sections, src_path, img_dir, n_images, output_dir, split_level="auto",
              min_size=MIN_CHAPTER_SECTIONS) -> dict:
    title = next(
        (s["text"] for s in sections if s.get("type") == "heading" and s.get("level") == 1),
        Path(src_path).stem,
    )
    chapter_level = _pick_split_level(sections, title, split_level)
    chapters = [] if chapter_level == 0 else _headings_at(sections, chapter_level, title)

    from collections import Counter
    type_counts = dict(Counter(s["type"] for s in sections))

    manifest = {
        "title": title,
        "source_file": str(src_path),
        "image_dir": str(img_dir),
        "total_images": n_images,
        "total_sections": len(sections),
        "chapter_level": chapter_level,
        "chapters": chapters,
        "type_counts": type_counts,
        "sections": sections,
    }
    out = Path(output_dir)
    (out / "manifest.json").write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")

    # Per-chapter JSON files for chunked, timeout-safe processing
    chapter_files = _write_chapter_files(sections, chapter_level, title, out, min_size)

    # Print summary
    print(f"✓ Title        : {title}")
    print(f"✓ Sections     : {len(sections)}  {type_counts}")
    if chapter_level == 0:
        print(f"✓ 模式         : --no-split（单文件，章节不拆分）")
    else:
        print(f"✓ Chapters (H{chapter_level}): {len(chapters)}")
    n_small = sum(1 for s in sections if s.get("type") == "image" and s.get("small_inline"))
    if n_small:
        print(f"✓ Images       : {n_images} → {img_dir}  (其中 {n_small} 张小图疑似公式/局部标签)")
    else:
        print(f"✓ Images       : {n_images} → {img_dir}")
    print(f"✓ Manifest     : {out / 'manifest.json'}")
    print(f"✓ Chapter files: {len(chapter_files)} → {out}/chapter_NN.json")
    print()
    if chapter_level == 0:
        print("[no-split] 全部 section 写入 chapter_01.json。")
        print("  写完 MD 后用 wc -c 检查文件大小；超过 5 MB 再按 H2 拆分为多文件。")
    elif chapters:
        print(f"章节列表 (H{chapter_level}):")
        for i, c in enumerate(chapters, 1):
            print(f"  {i:2d}. {c}")
        big = _max_chunk_size(sections, chapter_level, title)
        if big > MAX_CHAPTER_SECTIONS:
            print(f"\n[提示] 最大章节约 {big} 个 section，属大文件。写入时务必"
                  f"骨架→逐 ### 节 Edit 填充（120s 超时限制单次 Edit，不是文件大小）；"
                  f"若该章过于庞大想进一步细分，可对本文档加 --split-level 3。")
    else:
        print("[WARN] 未检测到分章节标题，将作为单文件处理。")
    return manifest


def _write_chapter_files(sections, chapter_level, title, out: Path,
                         min_size=MIN_CHAPTER_SECTIONS) -> list:
    """Split sections at each chapter-level heading; write chapter_NN.json each.

    Each chunk records `parent` — the nearest heading shallower than the split level —
    so the skill can build breadcrumbs / cross-links and name files hierarchically.

    Content conservation is guaranteed: every section ends up in exactly one chunk.
    A parent-level heading (and any preamble before the first child heading) is PREPENDED
    to the chunk it introduces — not appended to the previous one — and leading content
    before the first split heading rides along into the first chunk.

    chapter_level == 0 (--no-split): write all sections to a single chapter_01.json.
    """
    # --no-split mode: one file with everything
    if chapter_level == 0:
        fp = out / "chapter_01.json"
        fp.write_text(json.dumps({
            "index": 1,
            "heading": title,
            "parent": None,
            "headings": [title],
            "section_count": len(sections),
            "sections": sections,
        }, ensure_ascii=False, indent=2), encoding="utf-8")
        return [str(fp)]

    chapters = []
    current = None
    pending = []       # content to prepend to the next split-level chunk
    parent = None

    for s in sections:
        typ = s.get("type")
        lvl = s.get("level")
        text = s.get("text", "")
        is_split = (typ == "heading" and text != title and lvl == chapter_level)
        is_parent = (typ == "heading" and text != title
                     and lvl is not None and lvl < chapter_level)

        if is_parent:
            parent = text
            pending.append(s)           # parent heading leads the next chunk
            continue
        if is_split:
            if current:
                chapters.append(current)
            current = {"heading": text, "parent": parent, "sections": pending + [s]}
            pending = []
            continue
        # normal content (or the document title)
        if current is not None:
            current["sections"].append(s)
        else:
            pending.append(s)           # before the first split heading → buffer it

    if current:
        chapters.append(current)
    if pending:
        # Leftover: a parent section with no split-level child, or a doc with no split
        # headings at all. Emit as its own chunk so nothing is lost.
        head = next((x["text"] for x in pending if x.get("type") == "heading"),
                    parent or title)
        chapters.append({"heading": head, "parent": parent, "sections": pending})

    chapters = _merge_small_chunks(chapters, min_size, MAX_CHAPTER_SECTIONS)

    files = []
    for i, ch in enumerate(chapters, 1):
        fp = out / f"chapter_{i:02d}.json"
        fp.write_text(json.dumps({
            "index": i,
            "heading": ch["heading"],
            "parent": ch["parent"],
            "headings": ch.get("headings", [ch["heading"]]),
            "section_count": len(ch["sections"]),
            "sections": ch["sections"],
        }, ensure_ascii=False, indent=2), encoding="utf-8")
        files.append(str(fp))
    return files


def _merge_small_chunks(chunks, min_size, max_size):
    """Merge consecutive under-`min_size` chunks that share a parent, without exceeding
    `max_size`. Keeps a `headings` list of every split-level title folded into a file so
    the skill can name it (e.g. by the parent + a 1.1–1.4 range)."""
    if min_size <= 0 or not chunks:
        for ch in chunks:
            ch.setdefault("headings", [ch["heading"]])
        return chunks

    merged = []
    buf = None
    for ch in chunks:
        ch = dict(ch)
        ch.setdefault("headings", [ch["heading"]])
        if buf is None:
            buf = ch
            continue
        same_parent = ch["parent"] == buf["parent"]
        buf_small = len(buf["sections"]) < min_size
        fits = len(buf["sections"]) + len(ch["sections"]) <= max_size
        if same_parent and buf_small and fits:
            buf["sections"] += ch["sections"]
            buf["headings"] += ch["headings"]
        else:
            merged.append(buf)
            buf = ch
    if buf:
        merged.append(buf)
    return merged


# ---------------------------------------------------------------------------
# .doc → .docx auto-conversion (macOS textutil)
# ---------------------------------------------------------------------------

def convert_doc_to_docx(doc_path: str) -> str:
    """Convert legacy .doc to .docx via textutil (macOS). Returns new path."""
    if not shutil.which("textutil"):
        sys.exit("[ERROR] .doc format needs conversion. On macOS textutil is missing; "
                 "convert manually to .docx (e.g. open in Word → Save As .docx).")
    out_path = str(Path(doc_path).with_suffix(".docx"))
    out_path = os.path.join("/tmp", Path(out_path).name)
    try:
        subprocess.check_call(["textutil", "-convert", "docx", doc_path, "-output", out_path],
                              stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        print(f"[INFO] Converted .doc → {out_path}")
        return out_path
    except Exception as e:
        sys.exit(f"[ERROR] textutil conversion failed: {e}")


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Extract .docx/.doc/.pdf for doc-to-notes")
    parser.add_argument("doc_path", help="Absolute path to .docx / .doc / .pdf")
    parser.add_argument("--output-dir", default=None)
    parser.add_argument("--max-img-px", type=int, default=DEFAULT_MAX_IMG_PX,
                        help="Resize images larger than this on any side (default 2000)")
    parser.add_argument("--split-level", default="auto",
                        help="Heading level to split chapter files at: auto|2|3 (default auto)")
    parser.add_argument("--no-split", action="store_true",
                        help="Output all content in a single chapter_01.json; "
                             "no chapter splitting. After writing the MD check its size — "
                             "if > 5 MB, split manually by H2 headings.")
    parser.add_argument("--min-sections", type=int, default=MIN_CHAPTER_SECTIONS,
                        help=f"Merge small same-parent chapters below this size "
                             f"(default {MIN_CHAPTER_SECTIONS}; 0 disables merging)")
    args = parser.parse_args()

    doc_path = args.doc_path
    if not os.path.exists(doc_path):
        sys.exit(f"[ERROR] File not found: {doc_path}")

    if args.output_dir:
        output_dir = args.output_dir
    else:
        safe = re.sub(r"[^\w一-鿿.-]", "_", Path(doc_path).stem)
        output_dir = f"/tmp/doc_notes_{safe}"
    os.makedirs(output_dir, exist_ok=True)

    ext = Path(doc_path).suffix.lower()
    if ext == ".doc":
        doc_path = convert_doc_to_docx(doc_path)
        ext = ".docx"

    if args.no_split:
        split = 0
    else:
        split = args.split_level
        if split not in ("auto",):
            try:
                split = int(split)
            except ValueError:
                sys.exit("[ERROR] --split-level must be auto, 2, or 3")

    if ext == ".docx":
        extract_docx(doc_path, output_dir, args.max_img_px, split, args.min_sections)
    elif ext == ".pdf":
        extract_pdf(doc_path, output_dir, args.max_img_px, split, args.min_sections)
    else:
        sys.exit(f"[ERROR] Unsupported type: {ext}. Use .docx / .doc / .pdf")
